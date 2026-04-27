#!/usr/bin/env python3
"""运维手册 Markdown → Word 转换脚本"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elem)

def add_table_from_md(doc, header_line, rows):
    """从 markdown 表格行创建 Word 表格"""
    headers = [c.strip() for c in header_line.strip('|').split('|')]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2B579A')

    # 数据行
    for r_idx, row in enumerate(rows):
        cols = [c.strip() for c in row.strip('|').split('|')]
        for c_idx, val in enumerate(cols):
            if c_idx < len(headers):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                if r_idx % 2 == 1:
                    set_cell_shading(cell, 'F2F2F2')

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置各级标题样式
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '微软雅黑'
        hs.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.styles['Heading 1'].font.size = Pt(22)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(13)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 跳过分隔线和目录锚点
        if re.match(r'^---+\s*$', line) or re.match(r'^\{#.*\}$', line):
            i += 1
            continue

        # 跳过目录链接行（如 "1. [xxx](#xxx)"）
        if re.match(r'^\d+\.\s+\[.*\]\(#', line):
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,3})\s+(.+)', line)
        if m:
            level = len(m.group(1))
            text = re.sub(r'\{#.*\}', '', m.group(2)).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 表格
        if '|' in line and re.match(r'^\s*\|', line):
            # 收集表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i] and re.match(r'^\s*\|', lines[i]):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            if len(table_lines) >= 2:
                # 跳过分隔行（|---|---|）
                header = table_lines[0]
                data_rows = []
                for tl in table_lines[1:]:
                    if re.match(r'^\s*\|[\s\-:|]+\|', tl):
                        continue
                    data_rows.append(tl)
                if data_rows:
                    add_table_from_md(doc, header, data_rows)
                    doc.add_paragraph('')
            continue

        # 代码块
        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 代码段背景通过段落底纹实现
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F5F5F5',
            })
            pPr.append(shd)
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 引用块
        if line.startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            # 左侧边框
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '8',
                qn('w:color'): '2B579A',
            })
            pBdr.append(left)
            pPr.append(pBdr)
            # 解析文本中的加粗
            _add_inline_formatting(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1

    doc.save(docx_path)
    print(f'✅ 已生成: {docx_path}')


def _add_inline_formatting(paragraph, text):
    """解析行内加粗 **text** 和 `code`"""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    md_to_docx(
        '/home/lee/.openclaw/workspace/knowledge-base/docs/运维手册_v1.0.md',
        '/home/lee/.openclaw/workspace/knowledge-base/docs/RAG知识库管理系统_运维手册_v1.0.docx'
    )
